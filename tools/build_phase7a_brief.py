from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "11-release"
ASSET_DIR = OUT_DIR / "assets" / "phase-7a"
DOCX_PATH = OUT_DIR / "PHASE-7A-approved-development-brief.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(30, 35, 42)
MUTED = RGBColor(85, 85, 85)
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=INK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
        size = 16
        color = BLUE
    elif level == 2:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        size = 13
        color = BLUE
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        size = 12
        color = DARK
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = INK
    r.bold = bold
    r.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        r = p.add_run(item)
        r.font.size = Pt(10.5)
        r.font.color.rgb = INK


def add_callout(doc, title, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.width = Inches(6.3)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc, image_name, caption, width=6.2):
    path = ASSET_DIR / image_name
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        r = cap.add_run(caption)
        r.font.size = Pt(8.5)
        r.font.color.rgb = MUTED


def add_doc_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    if widths is None:
        widths = [Inches(6.3 / len(headers))] * len(headers)
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], LIGHT)
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=DARK, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_page_break(doc):
    doc.add_page_break()


def create_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    header = sec.header.paragraphs[0]
    header.text = "HRMS Platform | Phase 7A Approved Development Brief"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = MUTED
    footer = sec.footer.paragraphs[0]
    footer.text = "Approved Phase 7A scope | Generated 2026-06-29"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("PHASE 7A APPROVED DEVELOPMENT BRIEF")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("What we are going to build, who it serves, and how the approved HRMS platform will work")
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    add_image(
        doc,
        "team-meeting-clean.jpg",
        "Web image: Wikimedia Commons team meeting photo, used as a collaboration visual.",
        width=6.3,
    )
    add_callout(
        doc,
        "Executive conclusion",
        "Phase 7A is the first full development phase of the HRMS platform. It builds the secure SaaS foundation first, adds Branch/Office Hierarchy and Shift Foundation, then delivers Core HR, Leave, Attendance, Payroll with India compliance, and Standard Reports. The approved design is built for configuration, audit, tenant safety, branch-scoped operations, shift-aware time data, and future modules without changing core code.",
    )

    add_doc_table(
        doc,
        ["Approved Package", "Status"],
        [
            ["75 implementation documents", "Approved"],
            ["5 hardening standard/backlog documents", "Approved"],
            ["Total Phase 7A development documentation", "80/80 Approved"],
        ],
        [Inches(4.4), Inches(1.9)],
    )

    add_page_break(doc)

    add_heading(doc, "1. What Phase 7A Is", 1)
    add_para(
        doc,
        "Phase 7A is the foundation and first working product phase for the HRMS platform. It is not only a set of HR screens. It is the base that allows the product to become a configurable, multi-tenant, enterprise-grade HRMS.",
    )
    add_bullets(
        doc,
        [
            "It creates a secure tenant-aware platform where each customer is isolated.",
            "It allows one tenant to operate child branches/offices with separate branch admins and scoped employee access.",
            "It gives HR teams configurable rules, workflows, policies, reports, and approvals.",
            "It delivers everyday HR operations: employee records, self-service, branch assignment, leave, shift-aware attendance, payroll, payslips, and reports.",
            "It creates the architecture needed to add later modules such as documents, onboarding, service desk, recruitment, performance, learning, expenses, and AI without rebuilding the core.",
        ],
    )
    add_callout(
        doc,
        "Simple meaning",
        "Phase 7A is the platform proof: first build the strong base, then build the core HR workflows on top of that base.",
        fill="F7FAFC",
    )

    add_heading(doc, "2. Who This Phase Is For", 1)
    add_doc_table(
        doc,
        ["User group", "What they will use"],
        [
            ["Employees", "View profile, request leave, see attendance, access payslips, use self-service changes."],
            ["Managers", "Approve leave and attendance changes, view team data, track pending work, and see shift-aware attendance context."],
            ["HR Admins", "Manage employees, policies, leave setup, attendance rules, workflows, and reports."],
            ["Payroll Teams", "Run payroll, validate exceptions, approve results, publish payslips, manage India compliance."],
            ["Tenant Admins", "Manage configuration, branch/office hierarchy, feature flags, module settings, security, and approvals."],
            ["Auditors and Security Teams", "Review audit trails, time-machine history, access evidence, and export controls."],
            ["Leadership", "Get reliable standard reports and confidence that payroll and compliance are controlled."],
        ],
        [Inches(1.7), Inches(4.6)],
    )

    add_page_break(doc)

    add_heading(doc, "3. What We Are Going To Develop", 1)
    add_para(doc, "The approved scope is split into platform foundations and business modules.")
    add_doc_table(
        doc,
        ["Area", "Approved delivery"],
        [
            ["Tenant Catalog + RLS", "Customer isolation, tenant context, and row-level security."],
            ["Branch / Office Hierarchy", "Child branches/offices, branch admin scope, employee branch assignment, and branch-aware reporting."],
            ["Identity + RBAC/ABAC", "Login, roles, permissions, policy-based access, sessions, and emergency access controls."],
            ["Effective Dating", "History-aware records so the system knows what was true on any date."],
            ["Audit / Time Machine", "Every important action is traceable, searchable, and reconstructable."],
            ["Event Bus + Outbox", "Reliable module-to-module communication without tight coupling."],
            ["Rule Engine", "No hardcoded HR or payroll policies; rules are configurable, versioned, and testable."],
            ["Workflow Studio", "Approvals, delegation, SLA, escalation, versioned workflow definitions."],
            ["Configuration-as-Data", "Feature flags, forms, policies, rules, reports, navigation, and settings as governed data."],
            ["Core HR + ESS", "Employee master, assignments, org data, and self-service requests."],
            ["Leave", "Leave types, accruals, balances, approvals, calendars, and payroll impact."],
            ["Attendance + First Connector", "Punch capture, regularization, one connector framework, payroll attendance feed."],
            ["Shift Foundation", "Basic shift templates, employee shift assignment, overrides, shift resolver, and shift-aware attendance/payroll."],
            ["Payroll + India Compliance", "Salary structures, payroll runs, payslips, PF/ESI/PT/LWF/TDS/Form 16 foundation."],
            ["Standard Reports", "Operational and statutory reports with security, filters, exports, and audit."],
        ],
        [Inches(1.75), Inches(4.55)],
    )

    add_heading(doc, "4. What Is Not Being Built In Phase 7A", 1)
    add_bullets(
        doc,
        [
            "Recruitment, performance, LMS, engagement, expenses, assets, service desk, and advanced onboarding are later phases.",
            "Multi-country payroll packs are later; Phase 7A is India-first payroll and compliance foundation.",
            "Advanced roster planning, shift swaps, and workforce demand planning are later phases. Phase 7A includes the basic Shift Foundation needed for attendance and payroll.",
            "Advanced BI, marketplace, and full AI assistant capabilities are not part of the Phase 7A development scope.",
            "Customer-specific changes must not be hardcoded into core logic.",
        ],
    )

    add_page_break(doc)

    add_heading(doc, "5. How The Product Will Work For Users", 1)
    add_doc_table(
        doc,
        ["Workflow", "Simple operating flow"],
        [
            ["Employee self-service", "Employee requests a profile change or leave -> workflow routes approval -> result updates record and audit."],
            ["Leave", "Employee applies -> rules check balance/calendar -> manager approves -> balance ledger updates -> payroll/report events fire."],
            ["Branch operations", "Tenant admin creates offices -> branch admins work only within approved scope -> reports and exports respect that boundary."],
            ["Attendance", "Punches arrive from web/manual/connector -> effective shift is resolved -> daily summary calculated -> corrections approved -> payroll receives payable-day data."],
            ["Payroll", "Payroll team runs dry run from Core HR, Leave, and shift-aware Attendance -> resolves exceptions -> approval locks results -> payslips published -> reports generated."],
            ["Configuration", "Admin drafts a policy/workflow/rule -> system validates impact -> approval -> publish -> rollback available."],
            ["Audit and reports", "Authorized users search history, view as-of data, export evidence, and track who changed what."],
        ],
        [Inches(1.6), Inches(4.7)],
    )
    add_callout(
        doc,
        "Operating principle",
        "Every important change is approved where required, audited, evented, tenant-safe, and explainable.",
    )

    add_heading(doc, "6. Architecture In Simple Language", 1)
    add_image(
        doc,
        "server-racks-clean.jpg",
        "Web image: Wikimedia Commons server infrastructure photo, used as a platform architecture visual.",
        width=6.3,
    )
    add_bullets(
        doc,
        [
            "Multi-tenant SaaS: many customers can use the platform, but their data stays isolated.",
            "Branch-scoped SaaS: a tenant can have child branches/offices; tenant admins see all, while branch admins see only approved branches.",
            "Security first: every API checks tenant, user identity, role, permissions, and policy.",
            "Rules-as-data: HR and payroll policies live in configurable rule sets, not in code.",
            "Workflow-as-data: approval flows are versioned definitions, not hardcoded paths.",
            "Event-driven backbone: modules communicate through reliable events such as BranchScopeAssigned, LeaveApproved, ShiftDefinitionPublished, or PayrollRunPublished.",
            "Effective dating and time machine: changes are historical, explainable, and usable for payroll and reporting.",
            "Open for extension: new modules connect through manifests, APIs, events, providers, and configuration.",
        ],
    )

    add_page_break(doc)

    add_heading(doc, "7. Development Order", 1)
    add_para(doc, "The approved order is dependency-led. Business modules come only after the platform foundations they need.")
    add_doc_table(
        doc,
        ["Sprint", "Theme", "Expected result"],
        [
            ["S1", "Tenant + Identity + Branch Scope", "Every request is tenant-safe, branch-aware, and permission-aware."],
            ["S2", "Effective Dating + Audit + Events", "History, audit, and events work across the platform."],
            ["S3", "Rule + Workflow + Configuration", "Rules, approvals, and configuration are data-driven."],
            ["S4", "Core HR + ESS", "Employee master, branch assignment, and self-service are usable."],
            ["S5", "Leave", "Leave policies, applications, approvals, and balances work."],
            ["S6", "Attendance + Shift Foundation + Connector", "Shift-aware attendance capture and first connector feed payroll."],
            ["S7", "Payroll Engine", "Payroll run, components, calculations, and payslips work."],
            ["S8", "India Compliance", "PF, ESI, PT, LWF, TDS, Form 16 foundation, FBP, revisions."],
            ["S9", "Workflow Hardening", "Delegation, SLA, escalation, and migration are hardened."],
            ["S10", "Reports + Hardening", "Reports, exports, performance, security, and release readiness."],
        ],
        [Inches(0.7), Inches(1.8), Inches(3.8)],
    )

    add_heading(doc, "8. Security, Audit, and Quality Controls", 1)
    add_bullets(
        doc,
        [
            "Every table is tenant-scoped and every query must be tenant-filtered.",
            "Every module supports RBAC, ABAC, audit logging, and OpenAPI documentation.",
            "Critical operations need idempotency so retries do not duplicate payroll, leave, attendance, or workflow actions.",
            "PII and payroll-sensitive data require classification, masking, encryption, retention, and export controls.",
            "Tests include tenant-abuse cases, cross-tenant export attempts, role changes, impersonation audit, and security checks.",
            "Target coverage is at least 85%, with higher focus on payroll calculations, rules, ledger, and security paths.",
        ],
    )

    add_page_break(doc)

    add_heading(doc, "9. What Has Been Approved", 1)
    add_doc_table(
        doc,
        ["Document set", "Count", "Status"],
        [
            ["Business Requirements", "15", "Approved"],
            ["Technical Designs", "15", "Approved"],
            ["Database Designs", "15", "Approved"],
            ["UI Designs", "15", "Approved"],
            ["Test Plans", "15", "Approved"],
            ["Phase 7A Hardening Standards / Backlog", "5", "Approved"],
            ["Total Phase 7A development documentation", "80", "Approved"],
        ],
        [Inches(3.4), Inches(0.8), Inches(1.7)],
    )
    add_heading(doc, "10. How To Operate Phase 7A After Build", 1)
    add_bullets(
        doc,
        [
            "Admins configure tenants, branches/offices, roles, features, rules, workflows, calendars, shifts, payroll policies, and reports.",
            "Employees and managers use the portal for self-service, leave, attendance, approvals, and payslips.",
            "Payroll teams run dry runs first, resolve exceptions, seek approval, lock payroll, and publish payslips.",
            "Support and auditors use audit search, time-machine views, event trace, report history, and export evidence.",
            "Every change should be made through configuration, workflow, rule versions, or provider adapters, not direct code changes.",
        ],
    )

    add_page_break(doc)

    add_heading(doc, "11. Implementation Guardrails", 1)
    add_doc_table(
        doc,
        ["Guardrail", "Meaning"],
        [
            ["No code before approved docs", "This gate is now complete for Phase 7A documentation."],
            ["No hardcoded business rules", "Leave, attendance, payroll, workflow routing, and eligibility use Rule Engine/configuration."],
            ["No customer-specific core changes", "Customer behavior must be delivered through configuration, extensions, feature flags, providers, and plugins."],
            ["OpenAPI mandatory", "Every API must be versioned and documented before implementation is complete."],
            ["Event-driven where possible", "Approved events connect modules without direct database coupling."],
            ["Tenant isolation always", "Every module must enforce tenant context and prevent cross-tenant access."],
        ],
        [Inches(2.0), Inches(4.3)],
    )

    add_heading(doc, "12. Source and Image Credits", 1)
    add_bullets(
        doc,
        [
            "Approved internal documents: Phase 7A business requirements, technical designs, database designs, UI designs, test plans, and hardening standards under the project docs folder.",
            "Roadmap source: ROADMAP-001 phased delivery, approved roadmap.",
            "Web image: WLM international team meeting Vienna 2023-05-27 12, Wikimedia Commons, used as a collaboration visual.",
            "Web image: Wikimedia Servers-0051 19, Wikimedia Commons, used as an infrastructure visual.",
            "Web reference themes used in approved docs: Microsoft SQL Server temporal/RLS documentation, OpenAPI Specification, CloudEvents, RabbitMQ reliability, OpenFeature, OWASP API Security, WCAG 2.2, EPFO, ESIC, Income Tax India, and market references including Zoho People, Keka, greytHR, BambooHR, and Darwinbox.",
        ],
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    create_doc()
